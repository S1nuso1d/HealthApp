# -*- coding: utf-8 -*-
"""Генерация отчёта по практике HealthApp (~25 стр. A4) с диаграммами и макетами экранов."""

from __future__ import annotations

import os
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path(__file__).resolve().parent / "report_assets"
OUTPUT = Path(__file__).resolve().parent / "otchet_healthapp_praktika.docx"

ASSETS.mkdir(parents=True, exist_ok=True)

plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False


def save_fig(name: str, dpi: int = 150):
    path = ASSETS / name
    plt.tight_layout()
    plt.savefig(path, dpi=dpi, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


# --- Diagram generators ---

def diagram_architecture():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.set_title("Рисунок 1 — Архитектура системы HealthApp", fontsize=13, pad=12)

    def box(x, y, w, h, text, color):
        p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.15",
                           linewidth=1.5, edgecolor="#333", facecolor=color)
        ax.add_patch(p)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9, wrap=True)

    box(0.3, 3.8, 2.4, 1.2, "Android-клиент\n(Kotlin, Compose,\nHilt, Retrofit)", "#B8E6D8")
    box(3.5, 4.5, 2.8, 0.9, "REST API\n(FastAPI)", "#A8D4F0")
    box(3.5, 2.8, 2.8, 1.0, "Сервисный слой\n(AI, аналитика,\nумные напоминания)", "#D4E8F7")
    box(3.5, 1.0, 2.8, 1.0, "PostgreSQL / SQLite\n(SQLAlchemy ORM)", "#E8E8E8")
    box(7.0, 3.8, 2.5, 1.2, "Внешние источники\nHealth Connect,\nFatSecret, CSV", "#FFF3CD")
    box(7.0, 1.8, 2.5, 1.0, "WebSocket\n(RealtimeUpdates)", "#E2D5F5")

    for (x1, y1, x2, y2) in [(2.7, 4.4, 3.5, 4.9), (4.9, 4.5, 4.9, 3.8), (4.9, 2.8, 4.9, 2.0),
                               (6.3, 4.4, 7.0, 4.4), (6.3, 2.3, 7.0, 2.3)]:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#444", lw=1.5))
    save_fig("fig01_architecture.png")


def diagram_er():
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("Рисунок 2 — Логическая модель данных (основные сущности)", fontsize=13, pad=10)

    entities = [
        (0.5, 6.2, "users\nid, email, password"),
        (3.0, 6.2, "user_profiles\nцели, антропометрия"),
        (5.8, 6.2, "daily_health_summary\nагрегаты за день"),
        (0.5, 4.0, "sleep_records\nсон"),
        (3.0, 4.0, "meal_records\nпитание"),
        (5.8, 4.0, "hydration_records\nвода"),
        (8.3, 4.0, "activity_records\nактивность"),
        (0.5, 1.8, "health_samples\nпульс, вес, HC"),
        (3.0, 1.8, "user_states\nнастроение"),
        (5.8, 1.8, "action_plans\nплан действий"),
        (8.3, 1.8, "smart_reminders\nтриггеры"),
        (3.0, 0.2, "insights / analysis_runs\nаналитика и AI"),
    ]
    for x, y, t in entities:
        p = FancyBboxPatch((x, y), 2.2, 1.1, boxstyle="round,pad=0.02", linewidth=1.2,
                           edgecolor="#2E7D6B", facecolor="#E8F8F4")
        ax.add_patch(p)
        ax.text(x + 1.1, y + 0.55, t, ha="center", va="center", fontsize=7.5)
    ax.annotate("", xy=(3.0, 6.75), xytext=(2.7, 6.75), arrowprops=dict(arrowstyle="-|>", lw=1.2))
    ax.text(2.85, 7.0, "1:1", fontsize=8)
    for ux in [1.6, 4.1, 6.9]:
        ax.plot([1.6, ux], [6.2, 5.1], "k-", lw=0.8, alpha=0.5)
    ax.text(1.0, 5.5, "user_id →", fontsize=7, color="#555")
    save_fig("fig02_er_diagram.png")


def diagram_client_layers():
    fig, ax = plt.subplots(figsize=(9, 6))
    ax.axis("off")
    ax.set_title("Рисунок 3 — Слои Android-приложения (Clean Architecture)", fontsize=13, pad=10)
    layers = [
        ("Presentation (UI)", "Compose-экраны, ViewModel, UiState/Events", "#B8E6D8", 4.5),
        ("Domain", "Интерфейсы репозиториев, бизнес-правила", "#A8D4F0", 3.3),
        ("Data", "RepositoryImpl, Retrofit API, DTO, кэш", "#D4E8F7", 2.1),
        ("Core / DI", "Navigation, Theme, Hilt-модули, Constants", "#E8E8E8", 0.9),
    ]
    for title, desc, color, y in layers:
        p = FancyBboxPatch((1, y), 7, 0.95, boxstyle="round,pad=0.02", linewidth=1.3,
                           edgecolor="#333", facecolor=color)
        ax.add_patch(p)
        ax.text(1.2, y + 0.62, title, fontsize=11, fontweight="bold", va="center")
        ax.text(1.2, y + 0.28, desc, fontsize=8.5, va="center", color="#333")
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 6)
    save_fig("fig03_android_layers.png")


def diagram_auth_flow():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.set_title("Рисунок 4 — Блок-схема маршрутизации при запуске приложения", fontsize=12, pad=10)

    steps = [
        (5, 9.2, "Старт / SplashScreen"),
        (5, 8.0, "Есть JWT в TokenStorage?"),
        (2.5, 6.6, "Нет → Login / Register"),
        (7.5, 6.6, "Да → загрузка профиля"),
        (2.5, 5.2, "Гостевой режим?\n(демо без сервера)"),
        (7.5, 5.2, "Профиль заполнен?\n(возраст, рост, вес)"),
        (2.5, 3.8, "Dashboard\n(LocalDemoData)"),
        (7.5, 3.8, "Нет → Onboarding"),
        (7.5, 2.4, "Да → Dashboard\n+ Wellness API"),
        (5, 1.0, "Основная навигация\n(5 вкладок + стеки)"),
    ]
    for x, y, t in steps:
        w, h = (2.8, 0.75) if "?" not in t else (2.6, 0.9)
        if "?" in t:
            shape = mpatches.FancyBboxPatch((x - w/2, y - h/2), w, h, boxstyle="round,pad=0.02",
                                            linewidth=1.2, edgecolor="#C45C00", facecolor="#FFF4E6")
        else:
            shape = FancyBboxPatch((x - w/2, y - h/2), w, h, boxstyle="round,pad=0.02",
                                  linewidth=1.2, edgecolor="#2E6B8A", facecolor="#E8F4FC")
        ax.add_patch(shape)
        ax.text(x, y, t, ha="center", va="center", fontsize=8)
    arrow_pairs = [
        ((5, 8.55), (5, 8.45)), ((4.2, 7.6), (2.8, 6.95)), ((5.8, 7.6), (7.2, 6.95)),
        ((2.5, 6.15), (2.5, 5.65)), ((7.5, 6.15), (7.5, 5.65)), ((2.5, 4.75), (2.5, 4.15)),
        ((7.5, 4.75), (7.5, 4.15)), ((7.5, 3.35), (7.5, 2.85)), ((5, 2.0), (5, 1.45)),
    ]
    for start, end in arrow_pairs:
        ax.annotate("", xy=end, xytext=start,
                    arrowprops=dict(arrowstyle="->", color="#444", lw=1.2))
    ax.text(3.5, 7.2, "нет", fontsize=7)
    ax.text(6.2, 7.2, "да", fontsize=7)
    save_fig("fig04_auth_flow.png")


def diagram_dashboard_pipeline():
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.axis("off")
    ax.set_title("Рисунок 5 — Загрузка данных главного экрана (Dashboard)", fontsize=12, pad=8)
    nodes = ["DashboardViewModel", "WellnessRepository", "GET /dashboard/home", "DashboardCache",
             "UI: метрики, brief, mood, plan"]
    xs = [0.5, 2.5, 4.8, 7.0, 8.8]
    for i, (x, n) in enumerate(zip(xs, nodes)):
        p = FancyBboxPatch((x, 1.5), 1.7, 1.0, boxstyle="round,pad=0.02", linewidth=1.2,
                           facecolor=["#B8E6D8", "#A8D4F0", "#FFE082", "#E0E0E0", "#C8E6C9"][i],
                           edgecolor="#333")
        ax.add_patch(p)
        ax.text(x + 0.85, 2.0, n, ha="center", va="center", fontsize=7)
        if i < len(xs) - 1:
            ax.annotate("", xy=(xs[i+1], 2.0), xytext=(x + 1.7, 2.0),
                        arrowprops=dict(arrowstyle="->", lw=1.3))
    ax.text(5.5, 0.7, "при ошибке сети → офлайн-кэш", fontsize=8, color="#C62828", ha="center")
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 4)
    save_fig("fig05_dashboard_pipeline.png")


def mockup_screen(title: str, filename: str, sections: list[tuple[str, str]]):
    """Stylized UI mockup (mint/sky HealthApp style)."""
    fig, ax = plt.subplots(figsize=(4.2, 7.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    # phone frame
    frame = FancyBboxPatch((0.08, 0.02), 0.84, 0.96, boxstyle="round,pad=0.01,rounding_size=0.04",
                           linewidth=2, edgecolor="#222", facecolor="#FAFAFA")
    ax.add_patch(frame)
    # gradient header
    header = FancyBboxPatch((0.1, 0.78), 0.8, 0.18, boxstyle="round,pad=0.005",
                            linewidth=0, facecolor="#4EC9A8")
    ax.add_patch(header)
    ax.text(0.5, 0.87, title, ha="center", va="center", fontsize=11, color="white", fontweight="bold")
    y = 0.72
    for stitle, body in sections:
        card = FancyBboxPatch((0.12, y - 0.14), 0.76, 0.13, boxstyle="round,pad=0.01",
                              linewidth=0.8, edgecolor="#CCCCCC", facecolor="white")
        ax.add_patch(card)
        ax.text(0.16, y - 0.03, stitle, fontsize=8, fontweight="bold", va="top")
        ax.text(0.16, y - 0.09, body, fontsize=7, va="top", color="#555")
        y -= 0.16
    # bottom nav hint
    nav = FancyBboxPatch((0.1, 0.04), 0.8, 0.06, boxstyle="round,pad=0.005",
                         facecolor="#E8F5F0", edgecolor="#AAA")
    ax.add_patch(nav)
    ax.text(0.5, 0.07, "Главная · Сон · Еда · Профиль", ha="center", fontsize=6, color="#2E7D6B")
    save_fig(filename)


def generate_mockups():
    mockup_screen("HealthApp — Главная", "fig06_screen_dashboard.png", [
        ("Сводка дня", "Сон 7ч 20м · Вода 1.4 л · 6 240 шагов"),
        ("Wellness Score", "Общий балл 78 · серии привычек"),
        ("Настроение", "Быстрый check-in за сегодня"),
        ("AI Brief", "Краткая рекомендация на день"),
        ("План действий", "3 задачи · 1 выполнена"),
    ])
    mockup_screen("Вход и регистрация", "fig07_screen_auth.png", [
        ("Email / пароль", "JWT-сессия, восстановление пароля"),
        ("Регистрация", "Подтверждение email (код)"),
        ("Попробовать демо", "Гостевой режим без бэкенда"),
        ("Онбординг", "Цель, пол, рост, вес, активность"),
    ])
    mockup_screen("AI-ассистент", "fig08_screen_ai.png", [
        ("Чат", "Контекст: профиль + метрики дня"),
        ("Brief", "Ежедневная сводка от модели"),
        ("История", "Сохранённые ответы на сервере"),
    ])
    mockup_screen("Питание и активность", "fig09_screen_meal_activity.png", [
        ("Дневник питания", "КБЖУ, приёмы пищи, FatSecret"),
        ("Сканер штрихкода", "Быстрый ввод продукта"),
        ("Активность", "Шаги, тренировки, Health Connect"),
    ])
    mockup_screen("Профиль и настройки", "fig10_screen_profile.png", [
        ("Hero-аватар", "Цели, антропометрия, BMI"),
        ("Интеграции", "Health Connect, FatSecret"),
        ("Импорт CSV", "Пакетная загрузка health samples"),
        ("Экспорт отчёта", "Текстовый health-report share"),
    ])
    mockup_screen("Тёмная тема (Brutal)", "fig11_screen_brutal_theme.png", [
        ("Монохром", "Чёрный / серый / белый UI"),
        ("Виджет", "Glance: шаги и вода на рабочем столе"),
        ("Графики", "Canvas без @Composable внутри draw"),
    ])


def diagram_api_map():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.axis("off")
    ax.set_title("Рисунок 12 — Карта REST API бэкенда", fontsize=12, pad=8)
    groups = [
        ("/auth", "login, register, verify, password"),
        ("/profile", "CRUD профиля, аватар"),
        ("/dashboard", "home bundle, scores"),
        ("/sleep, /hydration, /meal, /activity", "дневники метрик"),
        ("/states", "настроение и самочувствие"),
        ("/action-plan", "планы и задачи"),
        ("/ai", "brief, chat, recommendations"),
        ("/analytics", "insights, analysis runs"),
        ("/smart", "reminders, triggers"),
        ("/import", "CSV/batch health samples"),
        ("/integrations", "FatSecret OAuth link"),
        ("/health", "vitals samples"),
        ("/ws", "realtime push"),
    ]
    y = 5.5
    for path, desc in groups:
        ax.add_patch(FancyBboxPatch((0.3, y), 9.4, 0.38, boxstyle="round,pad=0.01",
                                    facecolor="#E3F2FD", edgecolor="#1565C0", linewidth=0.8))
        ax.text(0.5, y + 0.19, path, fontsize=9, fontweight="bold", va="center")
        ax.text(3.2, y + 0.19, desc, fontsize=8, va="center", color="#333")
        y -= 0.42
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    save_fig("fig12_api_map.png")


def set_doc_defaults(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    # Ensure East Asia font for Word
    rfonts = style.element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def add_centered(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.first_line_indent = Cm(0)
    h.paragraph_format.line_spacing = 1.5
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    return p


def add_figure(doc, path: Path, caption: str, width_cm: float = 15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    for run in cap.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = True
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def add_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
    doc.add_paragraph()


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def expand_section_2(doc):
    add_heading(doc, "2.1. Выбор технологий и обоснование", 2)
    add_body(doc,
        "Для мобильного клиента выбран Kotlin как официальный язык Android с поддержкой "
        "корутин и нулевой безопасности типов, что снижает количество ошибок при асинхронной "
        "работе с сетью. Jetpack Compose заменяет XML-layout: декларативный UI упрощает "
        "создание анимированных карточек и градиентных заголовков, соответствующих дизайн-"
        "системе проекта. Hilt автоматизирует внедрение зависимостей ViewModel → Repository → Api."
    )
    add_body(doc,
        "Серверная часть на FastAPI обеспечивает высокую скорость разработки REST и "
        "автоматическую документацию Swagger UI, что ускоряет согласование контрактов с "
        "мобильной командой. SQLAlchemy 2.x даёт переносимость между SQLite (локальная "
        "разработка) и PostgreSQL (продакшен). Для AI используется внешний LLM-провайдер "
        "через сервисный слой: ключи не попадают в APK, запросы логируются и ограничиваются "
        "на стороне backend."
    )
    add_table(doc,
        ["Компонент", "Технология", "Назначение"],
        [
            ["UI", "Jetpack Compose", "Экраны, тема, навигация"],
            ["Сеть", "Retrofit + OkHttp", "REST, interceptors, JWT"],
            ["DI", "Hilt", "Singleton API, репозитории"],
            ["Состояние", "ViewModel + StateFlow", "MVVM, переживание поворота"],
            ["Локально", "DataStore, DashboardCache", "Токены, офлайн дашборд"],
            ["Виджет", "Glance", "Домашний экран Android"],
            ["Backend", "FastAPI + SQLAlchemy", "API и персистентность"],
            ["Realtime", "WebSocket", "Push-обновления клиента"],
        ],
    )
    add_heading(doc, "2.2. Взаимодействие компонентов при типичном сценарии", 2)
    add_body(doc,
        "Рассмотрим сценарий «Пользователь открыл приложение утром». SplashViewModel "
        "читает TokenStorage; при валидном JWT вызывается ProfileRepository.getProfile(). "
        "Если профиль неполный — маршрут onboarding; иначе — dashboard. DashboardViewModel "
        "в корутине вызывает WellnessRepository.loadHome(), который обращается к "
        "DashboardApi.getHome(). Ответ маппится в UiState: отдельные data class для "
        "метрик, brief, reminders, plan. Параллельно подписка на AppRefreshBus обновляет "
        "экран после записи воды или сна на другом табе. При ошибке HTTP 5xx/timeout "
        "показывается snackbar и подставляется кэш с пометкой «данные могут быть устаревшими»."
    )
    add_code_block(doc,
        "// Упрощённая цепочка (псевдокод)\n"
        "viewModelScope.launch {\n"
        "  _uiState.update { it.copy(loading = true) }\n"
        "  wellnessRepository.getDashboardHome()\n"
        "    .onSuccess { bundle -> applyHome(bundle) }\n"
        "    .onFailure { loadFromCache() }\n"
        "}"
    )


def expand_section_3(doc):
    add_heading(doc, "3.1. Детализация доменных сущностей", 2)
    entities_text = [
        ("SleepRecord", "Время отхода ко сну и пробуждения, длительность, качество, источник (ручной/Health Connect). Используется для расчёта дефицита относительно target_sleep_hours профиля."),
        ("MealRecord и SavedDish", "Приёмы пищи с КБЖУ; сохранённые блюда ускоряют повторный ввод. copy-day копирует рацион предыдущего дня — типичный UX для диетологического трекинга."),
        ("HydrationRecord", "Объём в мл и метка времени; агрегируется на дашборде в ProgressRing относительно target_water_ml."),
        ("ActivityRecord", "Шаги, дистанция, активные минуты; интеграция с Health Connect для автоматического импорта."),
        ("HealthSample", "Универсальная таблица для импорта CSV и показателей с датчиков (пульс, вес, SpO₂)."),
        ("UserState", "Субъективное настроение и заметки — вход для AI и аналитики паттернов."),
        ("ActionPlan / SmartReminder", "Поведенческий слой: задачи на день и проактивные напоминания на основе триггеров."),
        ("Insight / AnalysisRun", "Результаты аналитического движка и журнал запусков ML-конвейера."),
    ]
    for name, desc in entities_text:
        add_body(doc, f"{name}. {desc}")

    add_heading(doc, "3.2. Безопасность и конфиденциальность", 2)
    add_body(doc,
        "Пароли хранятся только в виде bcrypt-хэша. JWT имеет ограниченный срок жизни; "
        "клиент не логирует токен в release-сборке. Экран DataPrivacy описывает, какие "
        "данные отправляются на сервер и как запросить удаление. Импорт CSV выполняется "
        "по явному действию пользователя. OAuth FatSecret проходит через backend redirect — "
        "секреты приложения не вшиваются в мобильный клиент. Аватар ограничен 5 МБ "
        "(константа AVATAR_MAX_BYTES синхронизирована с backend)."
    )


def expand_section_4(doc):
    add_heading(doc, "4.1. Паттерн Repository и обработка ошибок", 2)
    add_body(doc,
        "Каждый домен (Auth, Profile, Meal, Sleep, Wellness…) имеет интерфейс в package "
        "domain.repository и реализацию data.repository.*Impl. Сетевые вызовы оборачиваются "
        "в runCatching или кастомный Result-адаптер; ошибки HTTP маппятся в UiText для "
        "локализованных сообщений. Такой подход отделяет «что показать пользователю» от "
        "«что вернул сервер» и упрощает unit-тесты ViewModel с фейковыми репозиториями."
    )
    add_heading(doc, "4.2. Каталог экранов приложения", 2)
    add_table(doc,
        ["Экран", "Назначение", "Ключевые компоненты"],
        [
            ["Dashboard", "Сводка дня", "DashboardHeroCard, WellnessCards, mood"],
            ["Sleep", "Учёт сна", "Графики Canvas, HC import"],
            ["Nutrition", "Дневник питания", "MealDiary, BarcodeScanner"],
            ["Hydration", "Вода", "ProgressRing, быстрые кнопки"],
            ["Activity", "Шаги и тренировки", "Charts, HC sync"],
            ["Profile", "Профиль и цели", "Hero avatar, AppCard секции"],
            ["AI Assistant", "Чат и brief", "AiAssistantViewModel"],
            ["Timeline", "Лента инсайтов", "Insights, analysis runs"],
            ["Action Plan", "Задачи", "CRUD через API"],
            ["Settings", "Импорт, интеграции", "CSV picker, OAuth link"],
        ],
    )


def expand_section_5(doc):
    add_heading(doc, "5.1. Модуль рекомендаций и уведомлений", 2)
    add_body(doc,
        "RecommendationsScreen отображает сохранённые рекомендации с фильтрацией; данные "
        "приходят с AI/аналитики и кэшируются на сервере в saved_recommendation. "
        "HealthNotificationHelper планирует локальные напоминания о воде и приёмах пищи; "
        "нажатие открывает нужный маршрут через Intent. Это дополняет smart_reminders "
        "серверной логики, которые генерируются по правилам и ML-триггерам."
    )
    add_heading(doc, "5.2. Health Connect и фоновая синхронизация", 2)
    add_body(doc,
        "HealthConnectReader запрашивает разрешения Android 14+ для чтения шагов, сна и "
        "пульса. HealthConnectForegroundSync оформлен как foreground service с уведомлением "
        "— требование платформы для длительного чтения. Прочитанные сэмплы конвертируются "
        "в HealthSampleCreateDto и отправляются пакетом на /import/batch, что исключает "
        "дублирование при повторном импорте за счёт идемпотентных ключей на backend."
    )
    add_heading(doc, "5.3. Realtime и обновление UI", 2)
    add_body(doc,
        "RealtimeUpdatesClient устанавливает WebSocket-соединение после авторизации. "
        "События (например, обновление dashboard или новый insight) публикуются в "
        "AppRefreshBus — подписчики ViewModel перезагружают только нужные секции, "
        "не сбрасывая scroll state всего экрана."
    )


def expand_section_6(doc):
    add_heading(doc, "6.1. Дизайн-система и переиспользуемые компоненты", 2)
    add_body(doc,
        "Ядро UI вынесено в core.ui.components: AppScreen (градиентный фон), ScreenHeader "
        "(иконка в скруглённом квадрате, subtitle), AppCard (радиус 26 dp, outline), "
        "AppButton, AppTextField, SectionHeader, Shimmer для skeleton-loading. "
        "Тема Material 3 в Theme.kt задаёт primaryContainer/secondaryContainer для "
        "индикаторов NavigationBar. AppBranding централизует градиенты и цвета графиков — "
        "это позволило внедрить брутальную тёмную тему без правки каждого экрана вручную."
    )
    add_body(doc,
        "Нижняя навигация использует полупрозрачный surface и NavigationBarItemDefaults "
        "с primaryContainer для выбранного пункта — соответствует каноническому UI-гайду. "
        "Иконки Extended Material (DirectionsWalk для активности, MonitorHeart на сплэше) "
        "поддерживают RTL через AutoMirrored."
    )


def expand_section_7(doc):
    add_heading(doc, "7.1. Матрица тестирования", 2)
    add_table(doc,
        ["№", "Сценарий", "Ожидаемый результат", "Статус"],
        [
            ["1", "Регистрация + код email", "JWT, переход в onboarding", "Выполнено"],
            ["2", "Вход с неверным паролем", "Сообщение об ошибке", "Выполнено"],
            ["3", "Гостевой режим", "Dashboard на LocalDemoData", "Выполнено"],
            ["4", "Запись воды 250 мл", "Обновление кольца прогресса", "Выполнено"],
            ["5", "Офлайн после загрузки dash", "Кэш DashboardCache", "Выполнено"],
            ["6", "AI-чат", "Ответ сервера в ленте", "Выполнено"],
            ["7", "Импорт CSV", "Записи в health_samples", "Выполнено"],
            ["8", "assembleDebug", "Успешная сборка JDK 17+", "Выполнено"],
        ],
    )
    add_heading(doc, "7.2. Ограничения и риски", 2)
    add_body(doc,
        "На этапе прототипа не реализованы автоматизированные UI-тесты Compose и нагрузочное "
        "тестирование API. Health Connect доступен не на всех устройствах. AI-ответы зависят "
        "от внешнего провайдера и требуют мониторинга квот. Для публикации в магазине "
        "необходимы privacy policy, подпись release-keystore и пентест API. Рекомендуется "
        "вынести BASE_URL в product flavors (dev/staging/prod)."
    )
    add_body(doc,
        "Несмотря на ограничения, проект демонстрирует полный цикл разработки "
        "прикладного ПО: от модели данных до визуально цельного мобильного продукта. "
        "Полученные артефакты — исходный код в репозитории HealthApp, Swagger-документация "
        "backend, настоящий отчёт с диаграммами и макетами экранов — могут быть представлены "
        "на защите практики без дополнительной доработки содержательной части."
    )


def build_document():
    diagram_architecture()
    diagram_er()
    diagram_client_layers()
    diagram_auth_flow()
    diagram_dashboard_pipeline()
    generate_mockups()
    diagram_api_map()

    doc = Document()
    set_doc_defaults(doc)

    # Title page
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", 12)
    add_centered(doc, "Федеральное государственное бюджетное образовательное учреждение", 12)
    add_centered(doc, "высшего образования", 12)
    add_centered(doc, "«НАЗВАНИЕ УНИВЕРСИТЕТА»", 12, bold=True)
    add_centered(doc, "(полное наименование, город)", 11)
    doc.add_paragraph()
    add_centered(doc, "Кафедра информационных технологий и программирования", 12)
    doc.add_paragraph()
    add_centered(doc, "ОТЧЁТ", 16, bold=True)
    add_centered(doc, "по производственной (технологической) практике", 14)
    doc.add_paragraph()
    add_centered(doc, "Тема: Разработка мобильной экосистемы персонального", 13)
    add_centered(doc, "мониторинга здоровья HealthApp", 13, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_centered(doc, "Выполнил: студент группы ___  _____________________", 12)
    add_centered(doc, "(подпись)", 10)
    doc.add_paragraph()
    add_centered(doc, "Руководитель от кафедры ___________  / _____________ /", 12)
    add_centered(doc, "(подпись)", 10)
    doc.add_paragraph()
    add_centered(doc, "Москва 2026", 12)
    add_page_break(doc)

    # TOC placeholder
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    toc_items = [
        "1. Введение и цели проекта",
        "2. Архитектура и технологический стек (2.1–2.2)",
        "3. Серверная часть: база данных и API (3.1–3.2)",
        "4. Мобильный клиент Android (4.1–4.2)",
        "5. Ключевой функционал и интеграции (5.1–5.3)",
        "6. Пользовательский интерфейс и результаты (6.1)",
        "7. Тестирование и выводы (7.1–7.2)",
        "Приложения",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0)
    add_page_break(doc)

    # Section 1
    add_heading(doc, "1. Введение и цели проекта", 1)
    add_body(doc,
        "Настоящий отчёт описывает инженерную работу над проектом HealthApp — "
        "мобильным приложением и серверным API для комплексного мониторинга здоровья. "
        "Раздел о компании-работодателе опущен по согласованию: материал сфокусирован "
        "на проектных результатах, архитектуре, базе данных, интерфейсе и тестировании."
    )
    add_body(doc,
        "В рамках производственной практики была спроектирована и реализована "
        "кросс-платформенная по логике, но ориентированная на Android экосистема "
        "HealthApp — приложение для сбора, хранения и интерпретации показателей "
        "здоровья пользователя: сон, гидратация, питание, физическая активность, "
        "биометрия и субъективное самочувствие. Практическая значимость проекта "
        "состоит в объединении разрозненных источников данных (ручной ввод, "
        "импорт CSV, Health Connect, внешние пищевые базы) в единую модель и "
        "предоставлении пользователю агрегированной «картины дня» на главном экране, "
        "дополненной аналитикой, планом действий и диалоговым AI-ассистентом."
    )
    add_body(doc,
        "Целью работы являлась разработка полнофункционального прототипа, "
        "пригодного для демонстрации на защите: клиент на Kotlin и Jetpack Compose "
        "с архитектурой MVVM + Clean Architecture, сервер на Python (FastAPI) с "
        "реляционным хранилищем и набором REST/WebSocket-контрактов. Задачи включали: "
        "(1) проектирование схемы данных и API; (2) реализацию модулей учёта метрик; "
        "(3) интеграцию AI и аналитических сервисов; (4) обеспечение устойчивости "
        "к сбоям сети за счёт локального кэша; (5) создание узнаваемого пользовательского "
        "интерфейса в фирменной палитре «мята + небо» с альтернативной тёмной "
        "бруталистской темой; (6) подготовку вспомогательных сценариев — гостевой режим, "
        "онбординг, домашний виджет, экспорт отчёта."
    )
    add_body(doc,
        "Объектом исследования выступают процессы персонального health-tracking "
        "и проектирование мобильных информационных систем. Предметом — методы "
        "построения клиент-серверных приложений с офлайн-поддержкой и интеграцией "
        "внешних SDK. В отчёте опущен традиционный раздел о профиле организации-"
        "практики: акцент сделан непосредственно на инженерных результатах проекта HealthApp."
    )
    add_heading(doc, "1.1. Актуальность и аналоги", 2)
    add_body(doc,
        "Рынок приложений для здоровья (mHealth) растёт за счёт носимых устройств и "
        "единых платформ вроде Apple Health и Google Health Connect. Пользователи ожидают "
        "не разрозненных счётчиков калорий, а целостной картины: сон влияет на активность, "
        "гидратация — на самочувствие, питание — на вес. HealthApp позиционируется как "
        "агрегатор с интерпретацией: серверная аналитика и AI формируют brief и план действий, "
        "а не только хранят сырые числа. Среди аналогов — MyFitnessPal (питание), Sleep Cycle "
        "(сон), Samsung Health (экосистема). Отличие нашего прототипа — открытый backend, "
        "единый дашборд wellness, гостевой режим для демонстрации и кастомизируемая "
        "визуальная айдентика для учебного проекта."
    )
    add_body(doc,
        "Методология работы включала итеративную разработку: сначала базовые CRUD метрик "
        "и авторизация, затем dashboard bundle, AI-модули, интеграции и полировка UI. "
        "Каждая итерация завершалась ручной проверкой на эмуляторе и Postman/Swagger. "
        "Такой подход соответствует agile-практикам и позволил к концу практики иметь "
        "работающий сквозной сценарий «регистрация → онбординг → запись данных → "
        "просмотр сводки → рекомендация AI»."
    )

    # Section 2
    add_heading(doc, "2. Архитектура и технологический стек", 1)
    add_body(doc,
        "Система построена по классической трёхзвенной схеме. Мобильный клиент "
        "обменивается JSON по HTTPS с backend-сервисом; тот, в свою очередь, "
        "использует ORM SQLAlchemy для доступа к СУБД и отдельные сервисные модули "
        "для AI-запросов, расчёта инсайтов и генерации умных напоминаний. "
        "Дополнительный канал WebSocket доставляет события обновления (например, "
        "после синхронизации данных), что позволяет обновлять UI без полного "
        "перезапуска экрана."
    )
    add_figure(doc, ASSETS / "fig01_architecture.png",
               "Рисунок 1 — Общая архитектура HealthApp (клиент — API — БД — внешние источники)")
    add_body(doc,
        "На стороне Android применён стек: Kotlin 2.x, Jetpack Compose, Navigation "
        "Compose, Hilt (DI), Retrofit + OkHttp, Coroutines/Flow, DataStore для токенов, "
        "Coil для изображений, Glance для виджета. Слои приложения разделены так, "
        "чтобы UI не зависел от деталей сети: ViewModel обращается к интерфейсам "
        "репозиториев domain-слоя, а реализации в data-слое инкапсулируют API и кэш."
    )
    add_figure(doc, ASSETS / "fig03_android_layers.png",
               "Рисунок 2 — Слои Android-приложения")
    add_body(doc,
        "Backend реализован на FastAPI с автогенерацией OpenAPI-схемы, Pydantic-моделями "
        "для валидации и JWT-аутентификацией. При старте приложения выполняется "
        "create_all для метаданных ORM и лёгкие миграции schema_patches — это упрощает "
        "развёртывание на машине разработчика без отдельного контура Alembic на этапе "
        "прототипа. CORS настроен в режиме, совместимом с мобильным клиентом."
    )
    expand_section_2(doc)

    add_page_break(doc)

    # Section 3
    add_heading(doc, "3. Серверная часть: база данных и API", 1)
    add_body(doc,
        "Логическая модель данных охватывает учётную запись пользователя, расширенный "
        "профиль с целями и нормами (калории, вода, шаги, сон), записи по доменам "
        "здоровья, а также сущности «второго контура» — состояния, инсайты, "
        "прогоны анализа, планы действий, умные напоминания и триггеры. Связь "
        "пользователя с профилем — один к одному; все дневные записи содержат "
        "внешний ключ user_id и временны́е метки для построения таймлайна."
    )
    add_figure(doc, ASSETS / "fig02_er_diagram.png",
               "Рисунок 3 — Основные сущности базы данных")
    add_body(doc,
        "Таблица users хранит email и хэш пароля; user_profiles — антропометрию, "
        "пол, цель (похудение, набор массы, сон, энергия), уровень активности и "
        "таргеты по макронутриентам. Отдельно ведутся health_samples — унифицированные "
        "измерения (пульс, вес, данные Health Connect) с типом и единицей измерения. "
        "daily_health_summary агрегирует показатели за календарный день для быстрой "
        "отдачи на dashboard/home без N+1 запросов по каждой метрике."
    )
    add_figure(doc, ASSETS / "fig12_api_map.png",
               "Рисунок 4 — Группы REST-эндпоинтов", width_cm=16)
    add_body(doc,
        "API спроектирован RESTful: ресурсы сна, воды, питания и активности поддерживают "
        "CRUD с фильтрацией по дате; модуль /import принимает пакетные JSON и CSV "
        "для миграции истории из других приложений; /integrations управляет OAuth-"
        "привязкой FatSecret. Модуль /ai предоставляет brief (краткая сводка), "
        "чат с контекстом профиля и сохранённые рекомендации. Analytics отдаёт "
        "insights и журнал analysis_runs — это основа экрана «Лента» на клиенте."
    )
    add_body(doc,
        "Пример фрагмента регистрации маршрутов в точке входа FastAPI:"
    )
    code = doc.add_paragraph()
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.left_indent = Cm(1)
    run = code.add_run(
        "app.include_router(auth_router)\n"
        "app.include_router(dashboard_router)\n"
        "app.include_router(ai_router)\n"
        "app.include_router(analytics_router)\n"
        "app.include_router(smart_router)\n"
        "app.include_router(data_import_router)"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    expand_section_3(doc)

    add_page_break(doc)

    # Section 4
    add_heading(doc, "4. Мобильный клиент Android", 1)
    add_body(doc,
        "Клиентское приложение HealtApp-front организовано по feature-модулям "
        "внутри одного app-модуля Gradle: auth, dashboard, profile, sleep, meal, "
        "hydration, activity, health (vitals), recommendations, timeline, actionplan, "
        "aicoach, settings (import, privacy, integrations, notifications), onboarding, widget. "
        "Навигация централизована в AppNavGraph: от Splash выполняется проверка токена "
        "и полноты профиля, после чего пользователь попадает на Dashboard или Onboarding."
    )
    add_figure(doc, ASSETS / "fig04_auth_flow.png",
               "Рисунок 5 — Блок-схема маршрутизации при запуске")
    add_body(doc,
        "Для сетевого слоя определены Retrofit-интерфейсы (AuthApi, ProfileApi, "
        "DashboardApi, WellnessRepository и др.), DTO зеркалят Pydantic-схемы backend. "
        "TokenStorage в DataStore сохраняет access token и флаг гостевого режима; "
        "интерцептор OkHttp подставляет заголовок Authorization. При отсутствии сети "
        "WellnessRepositoryImpl возвращает последний успешный ответ из DashboardCache — "
        "это критично для демонстрации на защите при нестабильном Wi‑Fi."
    )
    add_figure(doc, ASSETS / "fig05_dashboard_pipeline.png",
               "Рисунок 6 — Конвейер загрузки главного экрана")
    add_body(doc,
        "Нижняя навигация (AppBottomNavigation) закрепляет пять основных разделов; "
        "вторичные экраны (AI, план, лента, виталы, импорт) открываются поверх "
        "с сохранением back stack. Deep link из push-уведомлений обрабатывается "
        "через Intent extra EXTRA_NAV_ROUTE в MainActivity."
    )
    expand_section_4(doc)

    add_page_break(doc)

    # Section 5
    add_heading(doc, "5. Ключевой функционал и интеграции", 1)
    add_body(doc,
        "Функциональный объём HealthApp охватывает полный цикл данных о здоровье пользователя: "
        "сбор (ручной, автоматический, импорт), хранение на сервере, агрегация, визуализация, "
        "интерпретация (правила, аналитика, LLM) и обратная связь (план, напоминания, чат). "
        "Ниже описаны ключевые модули; детализация подсистем приведена в подразделах 5.1–5.3. "
        "Взаимосвязь модулей обеспечивается общим user_id и шиной событий AppRefreshBus на клиенте."
    )
    add_body(doc,
        "Главный экран (Dashboard) — центральный UX-узел. ViewModel параллельно "
        "загружает home bundle с сервера: метрики сна/воды/питания/активности, "
        "wellness scores, streaks, smart reminders, превью action plan, AI brief. "
        "Пользователь может отметить настроение (states API) без ухода с главной. "
        "Карточки оформлены компонентами DashboardSectionCard, DashboardWellnessCards "
        "с градиентными иконками согласно UI-гайду проекта."
    )
    add_figure(doc, ASSETS / "fig06_screen_dashboard.png",
               "Рисунок 7 — Макет главного экрана (сводка дня)", width_cm=7)
    add_body(doc,
        "Модуль аутентификации поддерживает вход, регистрацию с email-кодом, "
        "сброс пароля и кнопку «Попробовать демо» — гостевой режим с LocalDemoData "
        "для презентации без развёрнутого сервера. После регистрации запускается "
        "онбординг: сбор цели, пола, роста, веса и уровня активности с записью в профиль."
    )
    add_figure(doc, ASSETS / "fig07_screen_auth.png",
               "Рисунок 8 — Макет экранов входа и онбординга", width_cm=7)
    add_body(doc,
        "AI-ассистент (AiAssistantScreen) реализует чат с контекстом: профиль, "
        "сводка дня, последние метрики. Ответы приходят с backend; UI показывает "
        "состояния загрузки и ошибок. Экран «План действий» работает с реальным API "
        "(создание, отметка выполнения, удаление задач), mock-данные удалены. "
        "«Лента» (Timeline) объединяет insights, analysis runs и пользовательские states."
    )
    add_figure(doc, ASSETS / "fig08_screen_ai.png",
               "Рисунок 9 — Макет AI-ассистента", width_cm=7)
    add_body(doc,
        "Питание: дневник приёмов пищи, поиск, сохранённые блюда, копирование дня; "
        "интеграция FatSecret через OAuth link на backend; сканер штрихкода (CameraX). "
        "Активность и сон: ручной ввод + импорт Health Connect (HealthConnectReader, "
        "foreground sync service). Гидратация — быстрые пресеты объёма. Health vitals — "
        "графики Canvas с вынесением @Composable цветов за пределы draw-лямбды (требование "
        "компилятора Compose)."
    )
    add_figure(doc, ASSETS / "fig09_screen_meal_activity.png",
               "Рисунок 10 — Макет модулей питания и активности", width_cm=7)
    add_body(doc,
        "Настройки: импорт CSV (ImportViewModel, парсер ImportCsvFormat), экран "
        "конфиденциальности данных, интеграции, уведомления. Profile — hero с аватаром, "
        "две AppCard с целями и личными данными, ссылки на AI/ленту/план, экспорт "
        "HealthReportExporter (share текстового отчёта). Виджет HealthDashboardWidget "
        "на Glance показывает ключевые цифры на рабочем столе."
    )
    add_figure(doc, ASSETS / "fig10_screen_profile.png",
               "Рисунок 11 — Макет профиля и настроек", width_cm=7)
    expand_section_5(doc)

    add_page_break(doc)

    # Section 6
    add_heading(doc, "6. Пользовательский интерфейс и результаты", 1)
    add_body(doc,
        "Визуальная концепция зафиксирована в правилах проекта (healthapp-compose-ui): "
        "светлая тема — градиенты мята (#4EC9A8) и небо (#64B5F6), скругление карточек "
        "26 dp, ScreenHeader со «стеклянной» surface, SectionHeader с вертикальной "
        "градиентной полоской. Тёмная тема реализована в стиле digital brutalism: "
        "монохромная палитра, высокий контраст, отдельные хелперы AppBranding.kt "
        "(brandingGradient, heroBlockGradient, metricIconGradient) — светлая тема "
        "сохранена без изменений."
    )
    expand_section_6(doc)
    add_figure(doc, ASSETS / "fig11_screen_brutal_theme.png",
               "Рисунок 12 — Макет тёмной брутальной темы и виджета", width_cm=7)
    add_body(doc,
        "Результаты практики можно сформулировать количественно: реализовано более "
        "15 экранов Compose, 19 ORM-моделей на backend, 13+ групп API, офлайн-кэш "
        "дашборда, 2 темы оформления, виджет, guest mode, 5 доменов метрик здоровья. "
        "Качественный результат — целостный продуктовый прототип, пригодный для "
        "пользовательского тестирования и дальнейшего развития (Wear OS, семейный "
        "доступ, CI/CD со staging-окружением отложены в backlog)."
    )
    add_body(doc,
        "На рисунках 7–12 представлены стилизованные макеты экранов. При финальной "
        "вёрстке отчёта рекомендуется заменить их скриншотами с эмулятора Android Studio "
        "(Pixel 6, API 34) — достаточно сделать снимки Dashboard, Login, Profile, "
        "AiAssistant, Meal, Sleep в светлой и тёмной теме. Путь подстановки: "
        "Вставка → Рисунки → заменить файл в папке docs/report_assets/ с тем же именем."
    )

    add_page_break(doc)

    # Section 7
    add_heading(doc, "7. Тестирование и выводы", 1)
    add_body(doc,
        "Тестирование проводилось вручную по чек-листам сценариев: регистрация и вход, "
        "онбординг, заполнение метрик, синхронизация с backend, работа в гостевом режиме, "
        "отключение сети (проверка DashboardCache), импорт CSV, привязка FatSecret, "
        "чат AI, CRUD плана действий, экспорт отчёта, виджет после обновления данных. "
        "Сборка debug-варианта выполняется командой ./gradlew assembleDebug в Android Studio "
        "с JDK 17+. Backend запускается через uvicorn app.main:app. Выявленные на этапе "
        "интеграции ошибки компиляции (отсутствующие DTO, Unit-возвраты репозитория, "
        "Composable внутри Canvas, ColorProvider в Glance) устранены в ходе практики."
    )
    expand_section_7(doc)
    add_body(doc,
        "Выводы: поставленные цели достигнуты — создана связная система мониторинга "
        "здоровья с современным стеком и продуманным UX. Практика позволила закрепить "
        "навыки проектирования API, работы с ORM, построения многослойного Android-"
        "приложения, интеграции ML/AI-сервисов и проектирования отказоустойчивого "
        "клиента. Перспективы развития — публикация в Google Play, настройка облачного "
        "staging, автотесты (JUnit, MockWebServer, pytest), расширение Wear OS."
    )

    add_page_break(doc)
    add_heading(doc, "ПРИЛОЖЕНИЯ", 1)
    add_heading(doc, "Приложение А. Структура репозитория", 2)
    add_body(doc,
        "HealthApp/ — корень монорепозитория.\n"
        "• HealtApp-front/ — Android-клиент (app/src/main/java/.../features/*).\n"
        "• HealthApp-back/backend/ — FastAPI, app/models, app/api, app/services.\n"
        "• docs/ — отчёт и сгенерированные рисунки (report_assets/)."
    )
    add_heading(doc, "Приложение Б. Основные маршруты навигации", 2)
    routes = (
        "splash, login, register, forgot_password, onboarding, dashboard, profile, "
        "sleep, nutrition, hydration, activity, recommendations, timeline, action_plan, "
        "data_privacy, data_import, integrations, health_vitals, notifications, ai_assistant"
    )
    add_body(doc, f"NavRoute: {routes}.")
    add_heading(doc, "Приложение В. Список рисунков", 2)
    captions = [
        "Архитектура системы", "Модель данных", "Слои Android", "Маршрутизация при запуске",
        "Конвейер Dashboard", "Макет главного экрана", "Макет авторизации", "Макет AI",
        "Макет питания и активности", "Макет профиля", "Макет тёмной темы", "Карта REST API",
    ]
    for i, cap in enumerate(captions, 1):
        add_body(doc, f"Рисунок {i} — {cap}.")
    add_heading(doc, "Приложение Г. Инструкция по развёртыванию", 2)
    add_body(doc,
        "Backend: установить Python 3.11+, создать venv, pip install -r requirements.txt, "
        "задать переменные окружения (DATABASE_URL, SECRET_KEY, ключ AI-провайдера), "
        "запустить uvicorn app.main:app --reload --host 0.0.0.0 --port 8000. "
        "Android: открыть HealtApp-front в Android Studio, JDK 17, в local.properties "
        "или Constants указать BASE_URL машины в локальной сети (не localhost эмулятора — "
        "использовать 10.0.2.2 для AVD). Сборка: gradlew assembleDebug. "
        "Для демо без сервера — кнопка «Попробовать демо» на экране входа."
    )
    add_body(doc,
        "Для подстановки реальных скриншотов в отчёт: запустить приложение, сделать снимки "
        "экрана (Win+Shift+S / кнопка эмулятора), в Word щёлкнуть по макету рис. 7–12 → "
        "Заменить рисунок → выбрать PNG. Сохранить пропорции и подпись под рисунком."
    )

    doc.save(OUTPUT)
    # Копия на рабочий стол для удобства
    desktop_copy = Path.home() / "Desktop" / "otchet_healthapp_praktika.docx"
    try:
        import shutil
        shutil.copy2(OUTPUT, desktop_copy)
        print(f"Copy: {desktop_copy}")
    except OSError as e:
        print(f"Desktop copy skipped: {e}")
    chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"Saved: {OUTPUT} (~{chars} chars text)")
    print(f"Assets: {ASSETS}")


if __name__ == "__main__":
    build_document()
