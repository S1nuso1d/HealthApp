# -*- coding: utf-8 -*-
"""Сборка отчёта ВГТУ (преддипломная практика) — объёмные главы, листинги, правки по коду."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import generate_practice_report as gpr

DESKTOP_DOCX = Path.home() / "Desktop" / "Titulny_list_OTChETA.docx"
DESKTOP_DOC = Path.home() / "Desktop" / "Titulny_list_OTChETA.doc"
PROJECT_COPY = Path(__file__).resolve().parent / "Otchet_VGTU_HealthApp.docx"

LISTING_MAIN = '''app.include_router(auth_router)
app.include_router(profile_router)
app.include_router(sleep_router)
app.include_router(hydration_router)
app.include_router(meal_router)
app.include_router(activity_router)
app.include_router(states_router)
app.include_router(ai_router)
app.include_router(analytics_router)
app.include_router(smart_router)
app.include_router(action_plan_router)
app.include_router(dashboard_router)
app.include_router(ws_router)
app.include_router(data_import_router)
app.include_router(integrations_router)
app.include_router(health_router)'''

LISTING_AUTH = '''@router.post("/register/start", response_model=RegisterStartResponse)
def register_start(user_data: UserCreate, db: Session = Depends(get_db)):
    email = str(user_data.email).lower().strip()
    existing = db.query(User).filter(User.email == email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email уже зарегистрирован")
    # ... отправка кода на email, запись PendingRegistration'''

LISTING_RECOMMEND = '''@staticmethod
def generate_recommendations(db: Session, user_id: int) -> list[RecommendationItem]:
    insights = (
        db.query(Insight)
        .filter(Insight.user_id == user_id)
        .order_by(Insight.created_at.desc())
        .all()
    )
    for insight in insights:
        priority = PriorityScoring.score_priority(insight)
        recommendation = RecommendationBuilder.build_from_insight(insight, priority)
        # ... дедупликация и сортировка по приоритету'''

LISTING_WELLNESS = '''override suspend fun getDashboardHome(days: Int) = runCatching {
    dashboardApi.getHome(days).also { dashboardCache.save(it) }
}.recoverCatching {
    dashboardCache.load() ?: throw it
}'''

LISTING_NAV = '''data object Dashboard : NavRoute("dashboard")
data object Timeline : NavRoute("timeline")
data object ActionPlan : NavRoute("action_plan")
data object Integrations : NavRoute("integrations")
data object AiAssistant : NavRoute("ai_assistant")'''


def add_listing(doc, caption: str, code: str, lang: str = "Фрагмент"):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    cap = p.add_run(f"{caption}\n")
    cap.font.name = "Times New Roman"
    cap.font.size = Pt(12)
    cap.italic = True
    block = doc.add_paragraph()
    block.paragraph_format.first_line_indent = Cm(0)
    block.paragraph_format.left_indent = Cm(1)
    block.paragraph_format.line_spacing = 1.15
    run = block.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def title_page(doc: Document):
    for line, size, bold in [
        ("Министерство науки и высшего образования Российской Федерации", 12, False),
        ("Федеральное государственное бюджетное образовательное учреждение", 12, False),
        ("высшего образования", 12, False),
        ("«Воронежский государственный технический университет»", 12, True),
        ("(ФГБОУ ВО «ВГТУ»)", 11, False),
        ("факультет информационных технологий и компьютерной безопасности", 12, False),
        ("кафедра компьютерных интеллектуальных технологий проектирования", 12, False),
    ]:
        gpr.add_centered(doc, line, size, bold)
    doc.add_paragraph()
    gpr.add_centered(doc, "ОТЧЁТ ПО ПРЕДДИПЛОМНОЙ ПРАКТИКЕ", 16, True)
    doc.add_paragraph()
    gpr.add_centered(doc, "Тема: разработка мобильной экосистемы персонального", 13)
    gpr.add_centered(doc, "мониторинга здоровья HealthApp", 13, True)
    doc.add_paragraph()
    doc.add_paragraph()
    for line in [
        "Обучающийся _________________________   Д.О. Татаринцев",
        "Группа бПО-221    Вид практики: производственная (преддипломная)",
        "Наименование предприятия: ООО «Код Арт»",
        "Руководитель от кафедры: С.В. Игрунова",
        "Оценка _________________________",
        "",
        "Воронеж 2026",
    ]:
        gpr.add_centered(doc, line, 12)


def build():
    gpr.diagram_architecture()
    gpr.diagram_er()
    gpr.diagram_client_layers()
    gpr.diagram_auth_flow()
    gpr.diagram_dashboard_pipeline()
    gpr.generate_mockups()
    gpr.diagram_api_map()

    doc = Document()
    gpr.set_doc_defaults(doc)
    title_page(doc)
    gpr.add_page_break(doc)

    gpr.add_heading(doc, "СОДЕРЖАНИЕ", 1)
    for item in [
        "1. Введение",
        "2. Архитектура и технологический стек системы",
        "3. Серверная часть: база данных, API и безопасность",
        "4. Мобильное Android-приложение",
        "5. Аналитика, искусственный интеллект и интеграции данных",
        "6. Пользовательский интерфейс и результаты работы",
        "7. Тестирование, выводы и перспективы развития",
        "Список литературы",
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0)
    gpr.add_page_break(doc)

    # === 1. ВВЕДЕНИЕ (объёмная глава) ===
    gpr.add_heading(doc, "1. Введение", 1)
    gpr.add_body(doc,
        "В настоящее время цифровые технологии активно применяются в сфере здравоохранения "
        "и персонального мониторинга состояния организма. Пользователи используют мобильные "
        "приложения, фитнес-браслеты и платформу Health Connect для отслеживания сна, питания, "
        "гидратации и физической активности. При этом многие решения хранят данные разрозненно "
        "и не формируют единую аналитическую картину с персональными рекомендациями."
    )
    gpr.add_body(doc,
        "Актуальность разработки системы HealthApp обусловлена необходимостью объединить "
        "категории данных пользователя в единую клиент-серверную платформу: мобильный клиент "
        "на Android (Kotlin, Jetpack Compose), сервер на FastAPI и реляционное хранилище. "
        "Система не только фиксирует показатели, но и агрегирует их на главном экране (Dashboard), "
        "формирует insights, план действий, умные напоминания и поддерживает диалоговый "
        "AI-ассистент с учётом профиля и метрик дня."
    )
    gpr.add_body(doc,
        "В ходе преддипломной практики на базе ООО «Код Арт» была разработана информационная "
        "система HealthApp — мобильная экосистема анализа данных здоровья. Целью работы являлась "
        "разработка полнофункционального программного комплекса, включающего: мобильное "
        "Android-приложение; серверную часть; подсистему хранения данных; модуль аналитики "
        "на основе правил и insights; AI-модуль (краткий brief и чат с поддержкой LLM-провайдера, "
        "в том числе Ollama); интеграции Health Connect, FatSecret и импорт CSV."
    )
    gpr.add_body(doc,
        "Для достижения цели решены задачи: анализ предметной области; проектирование "
        "архитектуры; разработка структуры БД и REST API; реализация backend и Android-клиента; "
        "интеграция внешних источников данных; разработка UI в соответствии с дизайн-системой "
        "проекта; тестирование сквозных сценариев. Практическая значимость — возможность "
        "получать комплексную сводку состояния здоровья и рекомендации на основе совокупности факторов."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig01_architecture.png",
                   "Рисунок 1 – Общая архитектура системы HealthApp")

    # === 2. АРХИТЕКТУРА ===
    gpr.add_heading(doc, "2. Архитектура и технологический стек системы", 1)
    gpr.add_body(doc,
        "Система HealthApp построена по клиент-серверной схеме. Android-клиент обменивается "
        "данными с backend по REST API в формате JSON; дополнительно используется WebSocket "
        "для доставки событий обновления (RealtimeUpdatesClient на клиенте). Сервер выполняет "
        "валидацию запросов (Pydantic), аутентификацию (JWT), персистентность (SQLAlchemy), "
        "расчёт агрегатов и генерацию рекомендаций. Внешние источники — Health Connect, OAuth "
        "FatSecret, пакетный импорт health samples из CSV — подключаются через отдельные модули."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig02_er_diagram.png",
                   "Рисунок 2 – Логическая модель данных (основные сущности)", 16)
    gpr.add_table(doc,
        ["Компонент", "Технология", "Назначение в HealthApp"],
        [
            ["Mobile", "Kotlin, Jetpack Compose, MVVM, Hilt", "UI, ViewModel, DI"],
            ["Сеть", "Retrofit, OkHttp, Coroutines", "REST, JWT, офлайн-кэш"],
            ["Backend", "Python, FastAPI, Uvicorn", "API, бизнес-логика"],
            ["ORM / БД", "SQLAlchemy", "SQLite (разработка), PostgreSQL (продакшен)"],
            ["Аналитика", "RecommendationEngine, Insight", "Правила и приоритеты (не ML-pipeline)"],
            ["AI", "LLM-провайдер (Ollama и др.)", "Brief, чат, пояснение insights"],
            ["Realtime", "WebSocket", "Обновление UI без полной перезагрузки"],
        ],
    )
    gpr.add_body(doc,
        "Выбор FastAPI обусловлен автогенерацией OpenAPI (Swagger), удобством описания "
        "асинхронных обработчиков и быстрой итерацией контракта с мобильной командой. "
        "На клиенте Jetpack Compose позволил реализовать единую дизайн-систему: градиентные "
        "заголовки, карточки AppCard, нижнюю навигацию и альтернативную брутальную тёмную тему."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig03_android_layers.png",
                   "Рисунок 3 – Слои Android-приложения (Clean Architecture)")
    gpr.add_figure(doc, gpr.ASSETS / "fig04_auth_flow.png",
                   "Рисунок 4 – Маршрутизация при запуске (Splash, auth, onboarding, demo)")

    gpr.add_page_break(doc)

    # === 3. СЕРВЕР ===
    gpr.add_heading(doc, "3. Серверная часть: база данных, API и безопасность", 1)
    gpr.add_body(doc,
        "Backend расположен в каталоге HealthApp-back/backend. Структура проекта: api (маршруты), "
        "models (SQLAlchemy), schemas (Pydantic), services (email, аватары, FatSecret), db, core. "
        "При старте приложения выполняется create_all для метаданных ORM и лёгкие миграции "
        "schema_patches — это упрощает развёртывание учебного прототипа."
    )
    gpr.add_body(doc,
        "База данных хранит учётные записи (users), профили с целями и нормами (user_profiles), "
        "записи сна, гидратации, питания (meal_records), активности, унифицированные health_samples "
        "(импорт и датчики), субъективные состояния (user_states), insights и analysis_runs, "
        "планы действий (action_plan), умные напоминания (smart_reminders), сохранённые рекомендации. "
        "Связь пользователя с профилем — один к одному; все дневные сущности содержат user_id "
        "и временны́е метки для построения ленты Timeline на клиенте."
    )
    gpr.add_body(doc,
        "Реализованные группы REST API (префиксы роутеров в app/main.py): /auth, /profile, "
        "/sleep, /hydration, /meal (не /nutrition), /activity, /states, /dashboard, /action-plan, "
        "/analytics (в т.ч. /analytics/recommendations), /ai (brief, chat, recommendations), "
        "/smart, /import, /integrations, /health, WebSocket. Отдельного эндпоинта /timeline нет — "
        "лента на клиенте собирается из analytics и states."
    )
    add_listing(doc, "Листинг 1 – Подключение маршрутов FastAPI (фрагмент app/main.py)", LISTING_MAIN)
    gpr.add_figure(doc, gpr.ASSETS / "fig12_api_map.png",
                   "Рисунок 5 – Карта REST API backend")
    gpr.add_body(doc,
        "Авторизация: регистрация в два этапа (register/start — код на email, register/verify), "
        "вход по OAuth2PasswordRequestForm, смена и восстановление пароля. Пароли хранятся "
        "в виде bcrypt-хэша; доступ к маршрутам защищён зависимостью get_current_user. "
        "Секреты SMTP и LLM, ключи FatSecret — только на сервере (.env), не в APK."
    )
    add_listing(doc, "Листинг 2 – Начало регистрации с подтверждением email (фрагмент auth.py)",
                LISTING_AUTH)

    gpr.add_page_break(doc)

    # === 4. ANDROID ===
    gpr.add_heading(doc, "4. Мобильное Android-приложение", 1)
    gpr.add_body(doc,
        "Клиент HealtApp-front реализован на Kotlin с Jetpack Compose. Архитектура: MVVM, "
        "Clean Architecture, паттерн Repository. ViewModel обращается к интерфейсам domain-слоя; "
        "реализации в data-слое инкапсулируют Retrofit API, DTO и локальные хранилища (DataStore, "
        "DashboardCache). Hilt связывает зависимости. Навигация централизована в AppNavGraph."
    )
    gpr.add_body(doc,
        "Реализованные экраны и маршруты: авторизация, регистрация, восстановление пароля, "
        "онбординг, Dashboard, сон, питание (nutrition), гидратация, активность, рекомендации, "
        "Timeline, Action Plan, AI Assistant, профиль, виталы, импорт данных, конфиденциальность, "
        "интеграции, уведомления. Поддерживается гостевой режим «Попробовать демо» (LocalDemoData) "
        "для демонстрации без развёрнутого сервера. Deep link из push-уведомлений ведёт на нужный таб."
    )
    add_listing(doc, "Листинг 3 – Фрагмент маршрутов навигации (NavRoutes.kt)", LISTING_NAV)
    gpr.add_figure(doc, gpr.ASSETS / "fig05_dashboard_pipeline.png",
                   "Рисунок 6 – Загрузка Dashboard и офлайн-кэш")
    add_listing(doc,
                "Листинг 4 – Репозиторий wellness: сеть и кэш (WellnessRepositoryImpl.kt)",
                LISTING_WELLNESS)
    gpr.add_body(doc,
        "Сетевой слой: Retrofit-интерфейсы (DashboardApi, StatesApi, AiApi, …), interceptors "
        "для JWT, маппинг ошибок в UiText. При сбое getDashboardHome() клиент подставляет "
        "последний успешный ответ из DashboardCache — важно для нестабильной сети на защите."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig07_screen_auth.png",
                   "Рисунок 7 – Экран авторизации (макет; рекомендуется заменить скриншотом)", 7)
    gpr.add_figure(doc, gpr.ASSETS / "fig06_screen_dashboard.png",
                   "Рисунок 8 – Главный экран Dashboard", 7)

    gpr.add_page_break(doc)

    # === 5. АНАЛИТИКА + AI + ИНТЕГРАЦИИ ===
    gpr.add_heading(doc, "5. Аналитика, искусственный интеллект и интеграции данных", 1)
    gpr.add_body(doc,
        "Аналитический контур построен на сущностях Insight: серверные сервисы формируют "
        "наблюдения по метрикам пользователя; RecommendationEngine преобразует insights в "
        "рекомендации с приоритетом (PriorityScoring, RecommendationBuilder), дедупликацией "
        "и сортировкой. Это rule-based подход, а не обучение моделей scikit-learn на этапе "
        "прототипа (библиотеки pandas/numpy/sklearn указаны в requirements для возможного "
        "расширения, но в текущей версии app/ не задействованы)."
    )
    add_listing(doc, "Листинг 5 – Генерация рекомендаций из insights (RecommendationEngine.py)",
                LISTING_RECOMMEND)
    gpr.add_body(doc,
        "AI-модуль: эндпоинты /ai для ежедневного brief, чата с контекстом профиля и метрик, "
        "получения рекомендаций; провайдер LLM настраивается через LLM_PROVIDER (по умолчанию ollama). "
        "На клиенте AiAssistantScreen отображает диалог и состояния загрузки. Пример текста "
        "рекомендации в отчёте: «За последние 7 дней наблюдается недостаток сна…» — соответствует "
        "формату ответа системы при наличии соответствующих insights."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig08_screen_ai.png",
                   "Рисунок 9 – AI-ассистент", 7)
    gpr.add_body(doc,
        "Интеграции внешних данных: (1) Health Connect — HealthConnectReader и foreground-сервис "
        "синхронизации, импорт шагов, сна и др. в health_samples через /import/batch; "
        "(2) FatSecret — OAuth через /integrations, поиск продуктов в дневнике питания; "
        "(3) CSV — ImportCsvFormat на клиенте и /import/csv на сервере; (4) WebSocket — "
        "события в AppRefreshBus для точечного обновления карточек Dashboard."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig09_screen_meal_activity.png",
                   "Рисунок 10 – Питание и активность", 7)

    gpr.add_page_break(doc)

    # === 6. UI ===
    gpr.add_heading(doc, "6. Пользовательский интерфейс и результаты работы", 1)
    gpr.add_body(doc,
        "Интерфейс выполнен на Material Design 3 с собственной дизайн-системой: светлая тема — "
        "палитра «мята + небо», градиентные hero-блоки (DashboardHeader, ProfileHeroBlock), "
        "карточки AppCard (скругление 26 dp, обводка), SectionHeader с вертикальной градиентной "
        "полоской. Тёмная тема — digital brutalism (монохром, AppBranding.kt). Переключение темы "
        "в профиле; виджет HealthDashboardWidget на Glance выводит сводку на рабочий стол."
    )
    gpr.add_body(doc,
        "Dashboard объединяет метрики дня, wellness score, streaks, mood check-in, AI brief, "
        "превью action plan и smart reminders — данные приходят пакетом GET /dashboard/home. "
        "Профиль: аватар, цели, антропометрия, BMI, ссылки на AI, Timeline, план, экспорт "
        "текстового health-report. Timeline отображает insights, analysis runs и user states. "
        "Дополнительно: сканер штрихкода для продуктов, локальные напоминания о воде и еде."
    )
    gpr.add_figure(doc, gpr.ASSETS / "fig10_screen_profile.png",
                   "Рисунок 11 – Профиль и настройки", 7)
    gpr.add_figure(doc, gpr.ASSETS / "fig11_screen_brutal_theme.png",
                   "Рисунок 12 – Тёмная тема и виджет", 7)
    gpr.add_body(doc,
        "Итоги реализации: более 15 экранов Compose; 13+ групп API; 19 ORM-моделей; офлайн-кэш; "
        "две визуальные темы; гостевой режим; интеграции Health Connect и FatSecret; отчёт "
        "по практике с диаграммами. Репозиторий опубликован на GitHub с тегом "
        "milestone/practice-2026-05-15 для фиксации версии."
    )

    gpr.add_page_break(doc)

    # === 7. ТЕСТИРОВАНИЕ ===
    gpr.add_heading(doc, "7. Тестирование, выводы и перспективы развития", 1)
    gpr.add_body(doc,
        "Тестирование проводилось вручную и с использованием Swagger UI, Postman и Android-"
        "клиента. На backend добавлены pytest-тесты (импорт CSV, meal CRUD, удаление аккаунта). "
        "Сборка debug: ./gradlew assembleDebug (JDK 17+). Проверены регистрация с кодом email, "
        "JWT-вход, гостевой режим, CRUD метрик, Dashboard с кэшем при отключении сети, AI-чат, "
        "импорт CSV, Health Connect (на поддерживаемых устройствах)."
    )
    gpr.add_table(doc,
        ["Сценарий", "Ожидаемый результат", "Статус"],
        [
            ["Регистрация + код email", "JWT, переход в onboarding", "Успешно"],
            ["Вход / неверный пароль", "Токен / сообщение об ошибке", "Успешно"],
            ["Гостевой режим", "Dashboard на демо-данных", "Успешно"],
            ["Dashboard + офлайн", "Кэш DashboardCache", "Успешно"],
            ["AI brief и чат", "Ответ сервера", "Успешно"],
            ["Импорт CSV", "Записи health_samples", "Успешно"],
            ["Action plan CRUD", "Синхронизация с API", "Успешно"],
        ],
    )
    gpr.add_body(doc,
        "Выводы: цель практики достигнута — создан программный комплекс HealthApp с современным "
        "стеком, согласованным API и целостным UX. Закреплены навыки backend- и мобильной "
        "разработки, проектирования БД, REST, интеграции LLM и внешних SDK."
    )
    gpr.add_body(doc,
        "Перспективы: публикация в Google Play; CI/CD и staging; автотесты Compose; Wear OS; "
        "расширение ML-прогнозирования при наличии размеченных данных; семейный доступ."
    )

    gpr.add_page_break(doc)
    gpr.add_heading(doc, "Список литературы", 1)
    refs = [
        "Документация FastAPI [Электронный ресурс]. – URL: https://fastapi.tiangolo.com",
        "Документация Kotlin [Электронный ресурс]. – URL: https://kotlinlang.org",
        "Jetpack Compose [Электронный ресурс]. – URL: https://developer.android.com/jetpack/compose",
        "Документация SQLAlchemy [Электронный ресурс]. – URL: https://www.sqlalchemy.org",
        "Health Connect [Электронный ресурс]. – URL: https://developer.android.com/health-and-fitness/guides/health-connect",
        "ГОСТ 7.32–2017. Отчёт о научно-исследовательской работе.",
        "Васильев А.Н. Python на примерах и задачах. – М.: ДМК Пресс.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph(f"{i}. {r}")
        p.paragraph_format.first_line_indent = Cm(0)

    doc.save(DESKTOP_DOCX)
    doc.save(PROJECT_COPY)
    print(f"DOCX: {DESKTOP_DOCX}")
    print(f"Copy: {PROJECT_COPY}")


if __name__ == "__main__":
    build()
